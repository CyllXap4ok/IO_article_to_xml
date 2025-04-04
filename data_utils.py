import re
from enum import Enum
import xml.etree.ElementTree as XMLT
from xml.etree.ElementTree import Element
from xml.dom import minidom
from unidecode import unidecode
from typing import Callable
from docx import Document
from docx.table import _Cell
from article import ArticleData
from common import Language, AuthorRole
from author import Author
from workplace import Workplace, WorkplaceFactory


class FileType(Enum):
    Article = 163411
    Review = 274521
    EssentialInfo = 315236


class DataExtractor:

    def extract_data(
        self,
        filepaths: dict[FileType, list[str]],
        article_data: ArticleData,
        progress_callback: Callable
    ):
        total_elements = sum(map(len, filepaths.values()))

        for file_type, paths in filepaths.items():
            if paths:
                for path in paths:
                    doc = Document(path)
                    if file_type == FileType.Article:
                        self.__extract_data_from_article_header_table(doc, article_data)
                        self.__extract_data_from_article_text(doc, article_data)
                    elif file_type == FileType.Review:
                        reviewer_name = self.__extract_name_from_review_path(path)
                        self.__extract_data_from_review(reviewer_name, doc, article_data)
                    progress_callback(100 / total_elements)

    def __extract_data_from_article_header_table(self, doc: Document, article_data: ArticleData):
        unique_cells: list[_Cell] = []
        # Разбиение ячеек таблица на уникальные, т.к. объединенные ячейки считываются как разные с одинаковой инфой
        for cell in doc.tables[0]._cells:
            if cell not in unique_cells: unique_cells.append(cell)

        """ 
            СТРУКТУРА unique_cells:
            ___________________________
            Индекс   |   Cодержание
            ___________________________
              3      |   авторы
              4      |   места работы
              5      |   абстракт
              7      |   ключевые слова
        """

        article_data[Language.ENG].abstract = unique_cells[5].text.replace("Abstract\n", "")
        article_data[Language.ENG].keywords = self.__extract_keywords(unique_cells[7])
        article_data.authors = self.__extract_authors(unique_cells[3], unique_cells[4])

    @staticmethod
    def __extract_keywords(keywords_cell: _Cell) -> list[str]:
        return [keyword.strip() for keyword in re.split(r',\s+(?![^()]*\))', keywords_cell.text[10:-1])]

    def __extract_authors(self, authors_cell: _Cell, workplaces_cell: _Cell) -> list[Author]:
        workplaces: dict[str, Workplace] = self.__extract_workplaces(workplaces_cell)
        authors: list[Author] = [Author()]
        author_indexes: list[str] = []

        for paragraph in authors_cell.paragraphs:
            is_workplace_index = False
            for run in paragraph.runs:
                if run.font.superscript and run.text.strip('\n '):
                    if not is_workplace_index: author_indexes.append('')
                    author_indexes[-1] += run.text
                    is_workplace_index = True
                elif run.text.strip():
                    if run.text.strip().startswith('and '):
                        authors[-1][Language.ENG].surname += run.text[4:]
                    else:
                        authors[-1][Language.ENG].surname += run.text

                    if run.text.__contains__('*'): authors[-1].role = AuthorRole.Corresponding
                    if run.text.__contains__(','): authors.append(Author())
                    is_workplace_index = False

        for i in range(0, len(authors)):
            authors[i][Language.ENG].initials, authors[i][Language.ENG].surname = (
                authors[i][Language.ENG].surname.strip(' ,*').rsplit(maxsplit=1))

            authors[i][Language.ENG].workplaces = [
                workplaces[unidecode(index.strip())]
                for index in author_indexes[i].split(',')
            ]

        return authors

    @staticmethod
    def __extract_workplaces(workplaces_cell: _Cell) -> dict[str, Workplace]:
        workplaces: dict[str, Workplace] = {}
        workplace_index = ''
        workplace_text = ''

        for paragraph in workplaces_cell.paragraphs:
            if paragraph.runs[0].font.superscript:
                if workplace_text: workplaces[workplace_index] = WorkplaceFactory.create(workplace_text.replace('\n', ' '))
                workplace_index = paragraph.runs[0].text.strip()
                workplace_text = ''.join([run.text for run in paragraph.runs[1:]])
            else:
                workplace_text += paragraph.text
        workplaces[unidecode(workplace_index)] = WorkplaceFactory.create(workplace_text)

        return workplaces

    @staticmethod
    def __extract_data_from_article_text(doc: Document, article_data: ArticleData):
        is_article_text = True
        is_funding_text = False

        # Проходимся по параграфам в поисках нужных абзацев
        for paragraph in doc.paragraphs:

            # Некоторые абзацы одного стиля почему-то делятся на разные стили, несмотря на то,
            # что раны у всех могут быть bold, поэтому проверяем дополнительно их. При этой проверке вылезают
            # пустые абзацы, поэтому исключаем их дополнительным условием
            if paragraph.style.font.bold or all(run.font.bold for run in paragraph.runs) and paragraph.text != "":

                # Находим начало и конец записи для каждого пункта по ключевым словам
                if paragraph.text == "Acknowledgements":
                    is_funding_text = True
                elif paragraph.text == "Corresponding author":
                    is_article_text = False
                    is_funding_text = False
                elif paragraph.text == "References":
                    break

                # Добавляем заголовок в article_text без цифр в начале и пропускаем последующий код цикла
                if is_article_text: article_data[Language.ENG].text += paragraph.text.lstrip("0123456789. ") + " "
                continue

            if is_article_text: article_data[Language.ENG].text += paragraph.text + " "
            if is_funding_text: article_data[Language.ENG].funding += paragraph.text + " "

    @staticmethod
    def __extract_name_from_review_path(review_path: str) -> str:
        file_name = review_path.split('/')[-1].removesuffix('.docx')  # Путь файла без расширений
        parts = file_name.split('_')

        """ 
             СТРУКТУРА parts:
            ____________________
            Индекс  |   Часть
            ____________________
              0     |   referee
              1     |   report
              2     |   1 or 2
              3     |   Фамилия
              4     |   Инициалы
              5     |   ...
        """

        reviewer_last_name = parts[3]
        reviewer_initials = '. '.join(parts[4]) + '.'

        return reviewer_last_name + ' ' + reviewer_initials

    @staticmethod
    def __extract_data_from_review(reviewer_name: str, doc: Document, article_data: ArticleData):
        surname, initials = reviewer_name.rsplit(maxsplit=1)
        review = ''
        is_review = False
        for paragraph in doc.paragraphs:
            if 'замечания для передачи авторам' in paragraph.text.lower():
                is_review = True
                continue
            elif 'дополнительные замечания для редактора' in paragraph.text.lower():
                break

            if is_review: review += paragraph.text + '\n'

        author = Author()
        author.role = AuthorRole.Reviewer
        author[Language.RUS].surname = surname
        author[Language.RUS].initials = initials
        author[Language.RUS].review = review.strip('\n ')
        article_data.authors.append(author)


class DataSaver:

    def save_data(self, saving_path: str, article_data: ArticleData):
        self.__save_to_docx(saving_path, article_data)
        self.__save_to_xml(saving_path, article_data)

    @staticmethod
    def __save_to_docx(saving_path: str, article_data: ArticleData):
        path = saving_path + '.docx'

    def __save_to_xml(self, saving_path: str, article_data: ArticleData):
        path = saving_path + '.xml'

        article = XMLT.Element('article')

        pages = XMLT.SubElement(article, 'pages')
        pages.text = ''

        self.__add_authors(article, article_data.authors)

        art_type = XMLT.SubElement(article, 'artType')
        art_type.text = "ОБЯЗАТЕЛЬНЫЙ элемент"

        lang_publ = XMLT.SubElement(article, "langPubl")
        lang_publ.text = "ВАРИАТИВНЫЙ элемент"

        for parent, tag, attr in [
            (XMLT.SubElement(article, 'artTitles'), 'artTitle', 'title'),
            (XMLT.SubElement(article, 'abstracts'), 'abstract', 'abstract'),
            (XMLT.SubElement(article, 'fundings'), 'funding', 'funding')
        ]:
            for lang in Language:
                if value := getattr(article_data[lang], attr):
                    XMLT.SubElement(parent, tag, lang=lang.name).text = value

        self.__create_readable_xml(article, path)

    @staticmethod
    def __create_lang_sub_elements(parent: Element, tag: str) -> dict[Language, Element]:
        return {
            lang: XMLT.SubElement(parent, tag, lang=lang.name)
            for lang in Language
        }

    def __add_authors(self, parent: Element, authors: [Author]):
        authors_el = XMLT.SubElement(parent, 'authors')
        for num, author in enumerate(authors, 1):
            author_el = XMLT.SubElement(authors_el, 'author', num=str(num))

            individual_info = self.__create_lang_sub_elements(author_el, 'individInfo')

            match author.role:
                case (AuthorRole.Reviewer):
                    XMLT.SubElement(author_el, 'role').text = author.role.value
                    XMLT.SubElement(individual_info[Language.RUS], 'comment').text = author[Language.RUS].review
                case (AuthorRole.Corresponding):
                    XMLT.SubElement(author_el, 'correspondent').text = author.role.value

            self.__add_author_lang_elements(individual_info, 'surname', author)
            self.__add_author_lang_elements(individual_info, 'initials', author)
            self.__add_organisations(individual_info, author)

    @staticmethod
    def __add_author_lang_elements(parents: dict[Language, Element], attr: str, author: Author):
        for lang, parent in parents.items():
            XMLT.SubElement(parent, attr).text = getattr(author[lang], attr)

    @staticmethod
    def __add_organisations(parents: dict[Language, Element], author: Author):
        for lang, parent in parents.items():
            if author[lang].workplaces:
                towns, countries, names = zip(*[
                    (workplace.town, workplace.country, workplace.name)
                    for workplace in author[lang].workplaces
                ])

                XMLT.SubElement(parent, 'town').text = '; '.join(towns)
                XMLT.SubElement(parent, 'country').text = '; '.join(countries)
                XMLT.SubElement(parent, 'orgName').text = '; '.join(names)

    @staticmethod
    def __create_readable_xml(article: Element, path: str):
        xml_string = XMLT.tostring(article, encoding="utf-8")
        readable_xml = minidom.parseString(xml_string).toprettyxml(indent='  ')
        with open(path, "w", encoding="utf-8") as f:
            f.write(readable_xml)