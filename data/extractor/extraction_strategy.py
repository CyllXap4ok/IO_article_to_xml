import os
import re

from abc import ABC, abstractmethod

from docx.text.paragraph import Paragraph
from unidecode import unidecode
from docx import Document
from docx.table import _Cell

from data.article import ArticleData
from data.author import Author
from data.workplace import Workplace
from data.enum_const import Language, AuthorRole, Code


class DataExtractionStrategy(ABC):

    @abstractmethod
    def extract_data(self, path: str, data_holder: ArticleData):
        doc = self.get_doc(path)

    @staticmethod
    def get_doc(file_path) -> Document:
        if not file_path.lower().endswith('.docx'):
            raise ValueError("Файл должен иметь расширение .docx")

        if not os.path.exists(file_path):
            raise FileNotFoundError("Файл не найден.")

        return Document(file_path)


class ArticleExtractionStrategy(DataExtractionStrategy):

    def extract_data(self, path: str, data_holder: ArticleData):
        doc = self.get_doc(path)
        self.__extract_table_data(doc, data_holder)
        self.__extract_text_data(doc, data_holder)

    def __extract_table_data(self, doc: Document, data_holder: ArticleData):
        # Разбиение ячеек таблица на уникальные, т.к. объединенные ячейки считываются как разные с одинаковой инфой
        unique_cells: list[_Cell] = []
        for cell in doc.tables[0]._cells:
            if cell not in unique_cells: unique_cells.append(cell)

        """ 
            СТРУКТУРА unique_cells:
            _________________________________
            Индекс   |   Cодержание
            _________________________________
                     |   сайт
                     |   год, том (выпуск), сраницы
                     |   DOI: 
              2      |   
                     |   received
                     |   accepted
            _________|_______________________
              3      |   авторы
              4      |   места работы
              5      |   абстракт
              7      |   ключевые слова
        """

        data_holder[Language.ENG].abstract = unique_cells[5].text.replace("Abstract\n", "")
        data_holder[Language.ENG].keywords = self.__extract_keywords(unique_cells[7])
        data_holder.pages = self.__extract_pages(unique_cells[2])
        data_holder.codes[Code.DOI] = self.__extract_DOI(unique_cells[2])
        data_holder.received_date = self.__extract_date(unique_cells[2], "Received")
        data_holder.accepted_date = self.__extract_date(unique_cells[2], "Accepted")
        data_holder.authors = self.__extract_authors(unique_cells[3], unique_cells[4])

    @staticmethod
    def __extract_pages(cell: _Cell) -> str:
        return cell.paragraphs[1].text.split(',')[-1].strip()

    @staticmethod
    def __extract_DOI(cell: _Cell) -> list[str]:
        for paragraph in cell.paragraphs:
            if paragraph.text.startswith("DOI:"):
                return [paragraph.text[4:].strip()]

    @staticmethod
    def __extract_date(cell: _Cell, keyword: str) -> str:
        for paragraph in cell.paragraphs:
            if paragraph.text.startswith(keyword):
                return paragraph.text[len(keyword):].strip(', ')

    def __extract_keywords(self, keywords_cell: _Cell) -> list[str]:

        formatted_text = self.__extract_formatted_keywords(keywords_cell)

        return [
            keyword.strip('. ')
            for keyword in re.split(r',\s+(?![^()]*\))', formatted_text)
        ]

    @staticmethod
    def __extract_formatted_keywords(cell: _Cell) -> str:
        """
            Извлекает ключевые слова из параграфа с сохранением форматирования
            (курсив, жирный и т.д.).

            Возвращает строку с HTML-подобными тегами (<i>, <b>, <sub>, <sup>).
        """
        formatted_text = []
        remaining_chars_to_skip = len('Key words:')

        for paragraph in cell.paragraphs:

            for run in paragraph.runs:
                run_text = run.text
                run_length = len(run_text)

                if remaining_chars_to_skip > 0:
                    if run_length <= remaining_chars_to_skip:
                        # Весь run — часть префикса, пропускаем
                        remaining_chars_to_skip -= run_length
                        continue
                    else:
                        # Часть run'а — префикс, часть — полезный текст
                        run_text = run_text[remaining_chars_to_skip:]
                        remaining_chars_to_skip = 0

                    # Применяем форматирование
                if run.font.italic:
                    run_text = f"<i>{run_text}</i>"
                if run.font.bold:
                    run_text = f"<b>{run_text}</b>"
                if run.font.subscript:
                    run_text = f"<sub>{run_text}</sub>"
                if run.font.superscript:
                    run_text = f"<sup>{run_text}</sup>"

                formatted_text.append(run_text)

        return "".join(formatted_text)

    def __extract_authors(self, authors_cell: _Cell, workplaces_cell: _Cell) -> list[Author]:
        workplaces: dict[str, Workplace] = self.__extract_workplaces(workplaces_cell)
        authors: list[Author] = [Author()]
        author_indexes: list[str] = []

        for paragraph in authors_cell.paragraphs:
            is_workplace_index = False
            for run in paragraph.runs:
                if run.font.superscript and run.text.strip('\n '):
                    if not is_workplace_index: author_indexes.append('')
                    author_indexes[-1] += run.text
                    is_workplace_index = True
                elif run.text.strip():
                    if run.text.strip().startswith('and '):
                        authors[-1][Language.ENG].surname += run.text[4:]
                    else:
                        authors[-1][Language.ENG].surname += run.text

                    if run.text.__contains__('*'): authors[-1].role = AuthorRole.Corresponding
                    if run.text.__contains__(','): authors.append(Author())
                    is_workplace_index = False

        for i in range(0, len(authors)):
            authors[i][Language.ENG].initials, authors[i][Language.ENG].surname = (
                authors[i][Language.ENG].surname.strip(' ,*').rsplit(maxsplit=1))

            authors[i][Language.ENG].workplaces = [
                workplaces[unidecode(index.strip())]
                for index in author_indexes[i].split(',')
            ]

        return authors

    @staticmethod
    def __extract_workplaces(workplaces_cell: _Cell) -> dict[str, Workplace]:
        workplaces: dict[str, Workplace] = {}
        workplace_index = ''
        workplace_text = ''

        for paragraph in workplaces_cell.paragraphs:
            if paragraph.runs[0].font.superscript:
                if workplace_text:
                    text = workplace_text.replace('\n', ' ')
                    workplaces[workplace_index] = Workplace.Builder().parse(text).build()

                workplace_index = paragraph.runs[0].text.strip()
                workplace_text = ''.join([run.text for run in paragraph.runs[1:]])
            else:
                workplace_text += paragraph.text

        text = workplace_text.replace('\n', ' ')
        workplaces[unidecode(workplace_index)] = Workplace.Builder().parse(text).build()

        return workplaces

    def __extract_text_data(self, doc: Document, data_holder: ArticleData):

        is_funding_text = False

        for paragraph in doc.paragraphs:

            # Некоторые абзацы одного стиля почему-то делятся на разные стили, несмотря на то,
            # что раны у всех могут быть bold, поэтому проверяем дополнительно их. При этой проверке вылезают
            # пустые абзацы, поэтому исключаем их дополнительным условием
            if self.__is_bold(paragraph) and self.__is_arial_font(paragraph):

                if paragraph.text.strip() == "Acknowledgements":
                    is_funding_text = True
                elif paragraph.text.strip() == "Corresponding author":
                    break

                if paragraph.text:
                    # Добавляем заголовок в ArticleData.text без цифр в начале и пропускаем последующий код цикла
                    data_holder[Language.ENG].text += paragraph.text.strip("0123456789. ") + " "
                    continue

            data_holder[Language.ENG].text += paragraph.text + " "

            if is_funding_text:
                data_holder[Language.ENG].funding += paragraph.text + " "

        data_holder[Language.ENG].text = data_holder[Language.ENG].text.strip()
        data_holder[Language.ENG].funding = data_holder[Language.ENG].funding.strip()

    @staticmethod
    def __is_bold(paragraph: Paragraph):
        return paragraph.style.font.bold or all(run.font.bold for run in paragraph.runs)

    @staticmethod
    def __is_arial_font(paragraph: Paragraph):
        for run in paragraph.runs:
            if run.font.name and run.font.name.lower() != "arial":
                return False

        return True


class ReviewExtractionStrategy(DataExtractionStrategy):

    def extract_data(self, path: str, data_holder: ArticleData):
        doc = self.get_doc(path)

        surname, initials = self.__extract_name_from_review_path(path).split(maxsplit=1)

        review = ''
        is_review = False

        for paragraph in doc.paragraphs:
            if 'замечания для передачи авторам' in paragraph.text.lower():
                is_review = True
                continue
            elif 'дополнительные замечания для редактора' in paragraph.text.lower():
                break

            if is_review: review += paragraph.text + '\n'

        author = Author()
        author.role = AuthorRole.Reviewer
        author[Language.RUS].surname = surname
        author[Language.RUS].initials = initials
        author[Language.RUS].review = review.strip('\n ')
        data_holder.authors.append(author)

    @staticmethod
    def __extract_name_from_review_path(review_path: str) -> str:
        file_name = review_path.split('/')[-1].removesuffix('.docx')  # Путь файла без расширений
        parts = file_name.split('_')

        """ 
             СТРУКТУРА parts:
            ____________________
            Индекс  |   Часть
            ____________________
              0     |   referee
              1     |   report
              2     |   1 or 2
              3     |   Фамилия
              4     |   Инициалы
              5     |   ...
        """

        reviewer_last_name = parts[3]
        reviewer_initials = '. '.join(parts[4]) + '.'

        return reviewer_last_name + ' ' + reviewer_initials
